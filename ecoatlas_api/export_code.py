# ecoatlas_api/export_code.py
"""
Script d'export du projet ecoatlas_api dans un fichier Word (.docx).

- Section 1 : arborescence complète du dossier ecoatlas_api
- Section 2 : contenu des fichiers texte/code (Python, config, etc.)

Usage :
    cd ecoatlas_api
    python export_code.py

Nécessite :
    pip install python-docx
"""

from pathlib import Path
from docx import Document

# === Réglages de base ===

# Dossier racine du projet (celui où se trouve ce script)
ROOT_DIR = Path(__file__).resolve().parent

# Nom du fichier Word généré
OUTPUT_DOCX = ROOT_DIR / "export_code_ecoatlas_api.docx"

# Dossiers à ignorer dans l'arborescence
IGNORE_DIRS = {
    ".git",
    "__pycache__",
    ".venv",
    "venv",
    ".idea",
    ".vscode",
    ".mypy_cache",
    ".pytest_cache",
    ".ruff_cache",
    ".tox",
    "dist",
    "build",
    "node_modules",
}

# Extensions de fichiers binaires ou inutiles à ignorer
BINARY_EXTS = {
    ".pyc",
    ".pyo",
    ".pyd",
    ".so",
    ".dll",
    ".exe",
    ".jpg",
    ".jpeg",
    ".png",
    ".gif",
    ".svg",
    ".ico",
    ".webp",
    ".bmp",
    ".zip",
    ".rar",
    ".7z",
    ".tar",
    ".gz",
    ".tgz",
    ".xz",
    ".db",
    ".sqlite",
    ".sqlite3",
    ".log",
}

# Extensions de fichiers texte que l'on veut inclure dans la section "Code source"
TEXT_EXTS = {
    ".py",
    ".txt",
    ".md",
    ".rst",
    ".ini",
    ".cfg",
    ".conf",
    ".env",
    ".yml",
    ".yaml",
    ".json",
    ".toml",
    ".sql",
}

# Taille max pour inclure le contenu d'un fichier dans le docx (en octets)
MAX_FILE_SIZE = 300 * 1024  # 300 Ko, comme pour l'autre export


# === Fonctions utilitaires ===

def is_ignored_dir(path: Path) -> bool:
    """Retourne True si le dossier doit être ignoré."""
    return path.name in IGNORE_DIRS


def is_binary_or_ignored_file(path: Path) -> bool:
    """Retourne True si le fichier est binaire ou à ignorer (par extension)."""
    return path.suffix.lower() in BINARY_EXTS


def is_text_file(path: Path) -> bool:
    """Retourne True si le fichier doit être inclus dans la section 'Code source'."""
    # On ignore explicitement le fichier de sortie et ce script lui-même
    if path.name == OUTPUT_DOCX.name:
        return False
    if path.name == Path(__file__).name:
        return False

    # Fichiers binaires ou indésirables
    if is_binary_or_ignored_file(path):
        return False

    # On inclut uniquement certaines extensions
    return path.suffix.lower() in TEXT_EXTS


def build_tree_lines(root: Path) -> list[str]:
    """
    Construit une liste de lignes représentant l'arborescence du projet,
    avec des icônes 📁 / 📄 et des indentations, à la manière de l'autre export.
    """
    lines: list[str] = []

    def walk(dir_path: Path, level: int = 0):
        # Tri des entrées : d'abord dossiers, puis fichiers, le tout trié par nom
        dirs = []
        files = []
        for entry in dir_path.iterdir():
            if entry.is_dir():
                if is_ignored_dir(entry):
                    continue
                dirs.append(entry)
            else:
                files.append(entry)

        dirs.sort(key=lambda p: p.name.lower())
        files.sort(key=lambda p: p.name.lower())

        indent = "    " * level

        for d in dirs:
            lines.append(f"{indent}📁 {d.name}")
            walk(d, level + 1)

        for f in files:
            lines.append(f"{indent}📄 {f.name}")

    # Ligne racine
    lines.append(f"📁 {root.name}")
    walk(root, level=1)

    return lines


def add_architecture_section(doc: Document):
    """Ajoute la section 1 : architecture du projet."""
    doc.add_heading("1. Architecture du projet ecoatlas_api", level=1)
    doc.add_paragraph("")  # petite ligne vide

    tree_lines = build_tree_lines(ROOT_DIR)
    for line in tree_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        # optionnel : police monospace pour l'arborescence
        run.font.name = "Consolas"


def add_code_section(doc: Document):
    """Ajoute la section 2 : code source (contenu des fichiers texte)."""
    doc.add_page_break()
    doc.add_heading("2. Code source", level=1)
    doc.add_paragraph("")

    # On parcourt tous les fichiers récursivement
    for path in sorted(ROOT_DIR.rglob("*")):
        if path.is_dir():
            continue

        # On ne re-inspecte pas ce script ou le docx de sortie
        if path.name == Path(__file__).name:
            continue
        if path.name == OUTPUT_DOCX.name:
            continue

        if not is_text_file(path):
            continue

        rel_path = path.relative_to(ROOT_DIR).as_posix()

        # Titre "Début du fichier"
        doc.add_paragraph(f"Début du fichier : {rel_path}")

        # Gestion des gros fichiers
        size = path.stat().st_size
        if size > MAX_FILE_SIZE:
            doc.add_paragraph(f"⚠️ Fichier ignoré (> {MAX_FILE_SIZE // 1024} Ko).")
        else:
            try:
                content = path.read_text(encoding="utf-8", errors="replace")
            except Exception as e:  # en cas de souci d'encodage/permissions
                doc.add_paragraph(f"⚠️ Impossible de lire le fichier : {e}")
            else:
                # On met tout le contenu dans un bloc
                p = doc.add_paragraph()
                run = p.add_run(content)
                run.font.name = "Consolas"

        # Fin du fichier
        doc.add_paragraph(f"Fin du fichier : {rel_path}")
        doc.add_paragraph("")  # ligne vide de séparation


def main():
    print(f"📂 Dossier racine : {ROOT_DIR}")
    print(f"📝 Fichier Word de sortie : {OUTPUT_DOCX}")

    doc = Document()

    # Titre global
    doc.add_heading("EcoAtlas API — Export du code", 0)
    doc.add_paragraph(
        "Export automatique de la structure et du code du projet backend ecoatlas_api "
        "pour documentation et partage."
    )

    # Section 1 : arborescence
    add_architecture_section(doc)

    # Section 2 : code source
    add_code_section(doc)

    # Sauvegarde
    doc.save(OUTPUT_DOCX)
    print("✅ Export terminé.")
    print(f"📁 Fichier généré : {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
